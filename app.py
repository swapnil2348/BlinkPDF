from flask import Flask, render_template, request, redirect, url_for, send_file, abort, jsonify
import os
import uuid
from werkzeug.utils import secure_filename

import fitz  # PyMuPDF
from PyPDF2 import PdfReader, PdfWriter
from PIL import Image
from docx import Document
from pptx import Presentation
from reportlab.pdfgen import canvas


app = Flask(__name__)

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "outputs")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024  # 100MB


# ------------------------------------------------
# FULL NON-AI TOOL LIST (33 TOOLS – COMPLETE)
# ------------------------------------------------

TOOLS = [
    {'slug': 'merge-pdf', 'title': 'Merge PDF', 'desc': 'Combine multiple PDF documents easily.'},
    {'slug': 'split-pdf', 'title': 'Split PDF', 'desc': 'Extract or split PDF pages.'},
    {'slug': 'compress-pdf', 'title': 'Compress PDF', 'desc': 'Reduce PDF size.'},
    {'slug': 'optimize-pdf', 'title': 'Optimize PDF', 'desc': 'Clean and optimize PDF.'},
    {'slug': 'rotate-pdf', 'title': 'Rotate PDF', 'desc': 'Rotate PDF pages.'},
    {'slug': 'watermark-pdf', 'title': 'Watermark PDF', 'desc': 'Add watermark to PDF.'},
    {'slug': 'number-pdf', 'title': 'Number Pages', 'desc': 'Add page numbers.'},
    {'slug': 'protect-pdf', 'title': 'Protect PDF', 'desc': 'Add password protection.'},
    {'slug': 'unlock-pdf', 'title': 'Unlock PDF', 'desc': 'Remove PDF password.'},
    {'slug': 'repair-pdf', 'title': 'Repair PDF', 'desc': 'Fix broken PDFs.'},
    {'slug': 'organize-pdf', 'title': 'Organize PDF', 'desc': 'Reorder and delete pages.'},
    {'slug': 'sign-pdf', 'title': 'Sign PDF', 'desc': 'Add digital signature.'},
    {'slug': 'annotate-pdf', 'title': 'Annotate PDF', 'desc': 'Highlight and comment.'},
    {'slug': 'redact-pdf', 'title': 'Redact PDF', 'desc': 'Remove sensitive info.'},

    {'slug': 'pdf-to-word', 'title': 'PDF to Word', 'desc': 'Convert to DOCX.'},
    {'slug': 'word-to-pdf', 'title': 'Word to PDF', 'desc': 'Convert DOCX to PDF.'},
    {'slug': 'pdf-to-image', 'title': 'PDF to Image', 'desc': 'Convert pages to images.'},
    {'slug': 'image-to-pdf', 'title': 'Image to PDF', 'desc': 'Make PDF from images.'},
    {'slug': 'pdf-to-excel', 'title': 'PDF to Excel', 'desc': 'Extract tables to XLSX.'},
    {'slug': 'excel-to-pdf', 'title': 'Excel to PDF', 'desc': 'Convert Excel to PDF.'},
    {'slug': 'pdf-to-powerpoint', 'title': 'PDF to PowerPoint', 'desc': 'Convert to PPTX.'},
    {'slug': 'powerpoint-to-pdf', 'title': 'PowerPoint to PDF', 'desc': 'Convert to PDF.'},

    {'slug': 'ocr-pdf', 'title': 'OCR PDF', 'desc': 'Make searchable.'},
    {'slug': 'extract-text', 'title': 'Extract Text', 'desc': 'Extract text.'},
    {'slug': 'extract-images', 'title': 'Extract Images', 'desc': 'Pull images.'},
    {'slug': 'deskew-pdf', 'title': 'Deskew PDF', 'desc': 'Auto-straighten.'},
    {'slug': 'crop-pdf', 'title': 'Crop PDF', 'desc': 'Crop pages.'},
    {'slug': 'resize-pdf', 'title': 'Resize PDF', 'desc': 'Change page size.'},
    {'slug': 'flatten-pdf', 'title': 'Flatten PDF', 'desc': 'Flatten layers.'},
    {'slug': 'metadata-editor', 'title': 'Metadata Editor', 'desc': 'Edit PDF metadata.'},
    {'slug': 'fill-forms', 'title': 'Fill Forms', 'desc': 'Fill PDF forms.'},
    {'slug': 'background-remover', 'title': 'Remove Background', 'desc': 'Clean background.'}
]

# ------------------------------------------------
# FULL AI TOOL LIST (6 TOOLS)
# ------------------------------------------------

AI_TOOLS = [
    {"slug": "ai-editor", "title": "AI PDF Editor", "url": "/ai/editor"},
    {"slug": "ai-summarizer", "title": "AI Summarizer", "url": "/ai/summarizer-page"},
    {"slug": "ai-chat", "title": "Chat with PDF", "url": "/ai/chat-page"},
    {"slug": "ai-translate", "title": "AI Translator", "url": "/ai/translate-page"},
    {"slug": "ai-table-extract", "title": "AI Table Extractor", "url": "/ai/table-page"},
    {"slug": "ai-rewrite", "title": "AI Rewrite PDF", "url": "/ai/rewrite-page"}
]

SLUG_TO_TOOL = {tool["slug"]: tool for tool in TOOLS}
AI_SLUGS = {tool["slug"]: tool for tool in AI_TOOLS}

# ----------------------
# BASIC PAGES
# ----------------------

@app.route("/")
def index():
    return render_template("index.html", tools=TOOLS)


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/contact")
def contact():
    return render_template("contact.html")


@app.route("/privacy")
def privacy():
    return render_template("privacy.html")


@app.route("/terms")
def terms():
    return render_template("terms.html")


# ----------------------
# TOOL DYNAMIC PAGE
# ----------------------

@app.route("/tool/<slug>")
def tool_page(slug):
    tool = next((t for t in TOOLS if t.get("slug") == slug), None)
    if not tool:
        return redirect(url_for("index"))

    return render_template("tool_page.html", tool=tool)


# ----------------------
# FILE PROCESSING
# ----------------------

@app.route("/process", methods=["POST"])
def process_file():

    tool_slug = request.form.get("tool")

    if tool_slug not in SLUG_TO_TOOL:
        return jsonify({"error": "Invalid tool"}), 400

    files = request.files.getlist("file")

    if not files or files[0].filename == "":
        return jsonify({"error": "No file uploaded"}), 400

    input_paths = []

    for file in files:
        filename = secure_filename(file.filename)
        unique_name = str(uuid.uuid4()) + "_" + filename
        input_path = os.path.join(app.config["UPLOAD_FOLDER"], unique_name)
        file.save(input_path)
        input_paths.append(input_path)

    try:
        if tool_slug == "merge-pdf":
            output_path = merge_pdf(input_paths)

        elif tool_slug == "split-pdf":
            output_path = split_pdf(input_paths[0])

        elif tool_slug == "compress-pdf":
            output_path = compress_pdf(input_paths[0])

        elif tool_slug == "pdf-to-jpg":
            output_path = pdf_to_jpg(input_paths[0])

        elif tool_slug == "jpg-to-pdf":
            output_path = jpg_to_pdf(input_paths)

        elif tool_slug == "pdf-to-word":
            output_path = pdf_to_word(input_paths[0])

        elif tool_slug == "word-to-pdf":
            output_path = word_to_pdf(input_paths[0])

        else:
            return redirect(url_for("tool_page", slug=tool_slug))

        return send_file(output_path, as_attachment=True)

    except Exception as e:
        print("ERROR:", str(e))
        return jsonify({"error": "Processing failed"}), 500


# ----------------------
# TOOL FUNCTIONS
# ----------------------

def merge_pdf(input_paths):
    output_path = os.path.join(OUTPUT_FOLDER, f"merged_{uuid.uuid4()}.pdf")

    writer = PdfWriter()

    for path in input_paths:
        reader = PdfReader(path)
        for page in reader.pages:
            writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)

    return output_path


def split_pdf(input_path):
    output_path = os.path.join(OUTPUT_FOLDER, f"split_{uuid.uuid4()}.pdf")

    reader = PdfReader(input_path)
    writer = PdfWriter()

    if len(reader.pages) > 0:
        writer.add_page(reader.pages[0])

    with open(output_path, "wb") as f:
        writer.write(f)

    return output_path


def compress_pdf(input_path):
    output_path = os.path.join(OUTPUT_FOLDER, f"compressed_{uuid.uuid4()}.pdf")

    doc = fitz.open(input_path)
    doc.save(output_path, garbage=4, deflate=True)

    return output_path


def pdf_to_jpg(input_path):
    output_folder = os.path.join(OUTPUT_FOLDER, f"images_{uuid.uuid4()}")
    os.makedirs(output_folder, exist_ok=True)

    doc = fitz.open(input_path)

    image_paths = []

    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=150)
        out_path = os.path.join(output_folder, f"page_{i+1}.jpg")
        pix.save(out_path)
        image_paths.append(out_path)

    return image_paths[0] if image_paths else None


def jpg_to_pdf(input_paths):
    output_path = os.path.join(OUTPUT_FOLDER, f"image_to_pdf_{uuid.uuid4()}.pdf")

    images = [Image.open(p).convert("RGB") for p in input_paths]

    images[0].save(output_path, save_all=True, append_images=images[1:])

    return output_path


def pdf_to_word(input_path):
    output_file = os.path.join(OUTPUT_FOLDER, f"converted_{uuid.uuid4()}.docx")

    doc = fitz.open(input_path)
    word = Document()

    for page in doc:
        text = page.get_text()
        if text.strip():
            word.add_paragraph(text)

    word.save(output_file)

    return output_file


def word_to_pdf(input_path):
    output_path = os.path.join(OUTPUT_FOLDER, f"word_to_pdf_{uuid.uuid4()}.pdf")

    doc = Document(input_path)
    c = canvas.Canvas(output_path)

    y = 800
    for para in doc.paragraphs:
        if y < 40:
            c.showPage()
            y = 800
        c.drawString(40, y, para.text)
        y -= 15

    c.save()

    return output_path


# ----------------------
# ERROR HANDLERS
# ----------------------

@app.errorhandler(404)
def not_found(e):
    return render_template("404.html"), 404


@app.errorhandler(500)
def server_error(e):
    return render_template("500.html"), 500


# ----------------------
# RUN LOCAL
# ----------------------

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
